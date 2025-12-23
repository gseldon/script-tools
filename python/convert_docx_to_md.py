#!/usr/bin/env python3
"""
Скрипт для преобразования таблиц из docx файлов в markdown формат.
"""
import os
from pathlib import Path
from docx import Document
from docx.document import Document as DocumentType


def table_to_markdown(table) -> str:
    """Преобразует таблицу из docx в markdown формат."""
    if not table.rows:
        return ""
    
    markdown_lines = []
    
    # Получаем заголовки из первой строки
    headers = []
    for cell in table.rows[0].cells:
        header_text = cell.text.strip().replace('\n', ' ')
        headers.append(header_text)
    
    # Добавляем заголовки
    markdown_lines.append("| " + " | ".join(headers) + " |")
    
    # Добавляем разделитель
    markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    
    # Добавляем строки данных
    for row in table.rows[1:]:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ')
            # Экранируем символы pipe в содержимом
            cell_text = cell_text.replace('|', '\\|')
            cells.append(cell_text)
        markdown_lines.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(markdown_lines)


def docx_to_markdown(docx_path: Path, output_path: Path = None) -> str:
    """Преобразует docx файл в markdown, извлекая таблицы."""
    doc = Document(docx_path)
    
    markdown_parts = []
    
    # Добавляем заголовок с именем файла
    file_name = docx_path.stem
    markdown_parts.append(f"# {file_name}\n")
    
    # Обрабатываем все элементы документа
    for element in doc.element.body:
        if element.tag.endswith('tbl'):  # Таблица
            # Находим соответствующую таблицу в документе
            for table in doc.tables:
                if table._element == element:
                    table_md = table_to_markdown(table)
                    if table_md:
                        markdown_parts.append(table_md)
                        markdown_parts.append("")  # Пустая строка после таблицы
                    break
        elif element.tag.endswith('p'):  # Параграф
            # Извлекаем текст из параграфа
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if para:
                text = para.text.strip()
                if text:
                    # Определяем стиль заголовка
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            level_num = int(level)
                            markdown_parts.append(f"{'#' * level_num} {text}\n")
                        except ValueError:
                            markdown_parts.append(f"{text}\n")
                    else:
                        markdown_parts.append(f"{text}\n")
    
    markdown_content = "\n".join(markdown_parts)
    
    # Сохраняем в файл, если указан путь
    if output_path:
        output_path.write_text(markdown_content, encoding='utf-8')
        print(f"✅ Преобразован: {docx_path.name} -> {output_path.name}")
    
    return markdown_content


def main():
    """Основная функция для обработки всех docx файлов в текущей директории."""
    script_dir = Path(__file__).parent
    docx_files = list(script_dir.glob("*.docx"))
    
    if not docx_files:
        print("❌ Не найдено docx файлов в директории docs")
        return
    
    print(f"Найдено {len(docx_files)} docx файлов для преобразования\n")
    
    for docx_path in docx_files:
        output_path = script_dir / f"{docx_path.stem}.md"
        try:
            docx_to_markdown(docx_path, output_path)
        except Exception as e:
            print(f"❌ Ошибка при обработке {docx_path.name}: {e}")


if __name__ == "__main__":
    main()

